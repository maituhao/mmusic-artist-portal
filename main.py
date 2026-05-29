import os
import json
import uuid
import shutil
import datetime
from fastapi import FastAPI, Request, Form, File, UploadFile
from fastapi.responses import JSONResponse, FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional

import firebase_admin
from firebase_admin import credentials, firestore
import drive_service

# --- Paths ---
BASE_DIR = os.path.dirname(__file__)
TEMPLATE_DIR = os.path.join(BASE_DIR, 'templates')

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.mount("/static", StaticFiles(directory=os.path.join(BASE_DIR, "static")), name="static")

# --- Firebase Init ---
firebase_creds_str = os.getenv("FIREBASE_CREDS_JSON")
try:
    if firebase_creds_str:
        cred_dict = json.loads(firebase_creds_str)
        cred = credentials.Certificate(cred_dict)
    else:
        cred = credentials.Certificate(os.path.join(BASE_DIR, "firebase_creds.json"))
    firebase_admin.initialize_app(cred)
    db = firestore.client()
except Exception as e:
    print("Warning: Firebase not initialized:", e)
    db = None

# --- Models ---
class TrackInfo(BaseModel):
    title: str
    version: str = "Original"
    mainArtist: str = ""
    featArtist: str = ""
    profile: str = ""
    composer: str = ""
    writer: str = ""
    producer: str = ""
    genre: str = ""
    lang: str = "Vietnamese"
    year: str = ""
    explicit: str = "No"
    tiktok: str = ""
    lyrics: str = ""

class PRFormRequest(BaseModel):
    q1: Optional[str] = ""
    q2: Optional[str] = ""
    q3: Optional[str] = ""
    q4: Optional[str] = ""
    q5: Optional[str] = ""
    q6: Optional[str] = ""
    q7: Optional[str] = ""
    q8: Optional[str] = ""
    q9: Optional[str] = ""
    q10: Optional[str] = ""
    q11: Optional[str] = ""
    q12: Optional[str] = ""
    q13: Optional[str] = ""

class LabelFormRequest(BaseModel):
    productTitle: str
    productType: str
    productArtist: str
    genre: str
    releaseDate: str
    releaseTime: str = ""
    preorderDate: str = ""
    preorderTime: str = ""
    tracks: List[TrackInfo]

# --- Helper functions ---
def get_release(release_id: str):
    if not db:
        return None
    doc = db.collection('releases').document(release_id).get()
    if doc.exists:
        return doc.to_dict()
    return None

def update_release_status(release_id: str, section_key: str):
    if not db:
        return
    doc_ref = db.collection('releases').document(release_id)
    doc = doc_ref.get()
    if doc.exists:
        data = doc.to_dict()
        if "status" not in data:
            data["status"] = {}
        if section_key not in data["status"]:
            data["status"][section_key] = {}
        data["status"][section_key]["completed"] = True
        data["status"][section_key]["timestamp"] = datetime.datetime.now().isoformat()
        doc_ref.set(data)

# --- Routes ---
@app.api_route("/", methods=["GET", "HEAD"])
async def root_redirect():
    from fastapi.responses import RedirectResponse
    return RedirectResponse(url="/portal")

@app.api_route("/portal", methods=["GET", "HEAD"], response_class=HTMLResponse)
async def serve_portal(id: Optional[str] = None):
    portal_path = os.path.join(BASE_DIR, "static", "portal.html")
    if not os.path.exists(portal_path):
        return HTMLResponse("File not found", status_code=404)
        
    with open(portal_path, 'r', encoding='utf-8') as f:
        html_content = f.read()

    title = "Artist Upload Portal | MMUSIC"
    meta_tags = ""

    if id:
        release_data = get_release(id)
        if release_data:
            artist_name = release_data.get("artist_name")
            product_name = release_data.get("product_name")
            if artist_name:
                title = f"Artist Upload Portal | {artist_name} | MMUSIC"
                meta_tags = f"""
    <meta property="og:title" content="Artist Upload Portal | {artist_name} | MMUSIC" />
    <meta property="og:description" content="Cổng tải lên tài nguyên sản phẩm {product_name or ''} của nghệ sĩ {artist_name}." />
    <meta property="og:type" content="website" />
    <meta property="og:image" content="/static/logo.png" />
                """

    html_content = html_content.replace(
        "<title>Artist Upload Portal | MMUSIC</title>",
        f"<title>{title}</title>{meta_tags}"
    )

    return HTMLResponse(content=html_content, status_code=200)

@app.get("/api/release/{release_id}")
async def get_release_status(release_id: str):
    data = get_release(release_id)
    if data:
        return JSONResponse({
            "artist_name": data.get("artist_name"),
            "product_name": data.get("product_name"),
            "status": data.get("status", {})
        })
    return JSONResponse({"error": "Không tìm thấy release"}, status_code=404)

@app.post("/api/release/{release_id}/label-form")
async def fill_label_form(release_id: str, payload: LabelFormRequest):
    release_data = get_release(release_id)
    if not release_data:
        return JSONResponse({"error": "Release not found"}, status_code=404)
        
    folder_ids = release_data.get("folder_ids", {})
    target_folder_id = folder_ids.get("1. Label form (track info)")
    
    if not target_folder_id:
        return JSONResponse({"error": "Label form folder not found"}, status_code=500)
    
    template_path = os.path.join(TEMPLATE_DIR, "[MMusic Records] Label Form - Product Info.xlsx")
    if not os.path.exists(template_path):
        return JSONResponse({"error": "Template file not found on server"}, status_code=500)
        
    try:
        import openpyxl
        wb = openpyxl.load_workbook(template_path)
        ws = wb.active
        
        ws['C4'] = payload.productTitle
        ws['C5'] = payload.productType
        ws['C6'] = payload.productArtist
        ws['C7'] = payload.genre
        ws['C8'] = payload.releaseDate
        if payload.releaseTime:
            ws['C9'] = payload.releaseTime
            
        preorder_value = ""
        if payload.preorderDate:
            preorder_value = payload.preorderDate
            if payload.preorderTime:
                preorder_value += " " + payload.preorderTime
        elif payload.preorderTime:
            preorder_value = payload.preorderTime
        ws['C10'] = preorder_value
        
        start_row = 16
        for i, track in enumerate(payload.tracks):
            row = start_row + i
            ws.cell(row=row, column=1, value=i+1)
            ws.cell(row=row, column=2, value=track.title)
            ws.cell(row=row, column=4, value=track.version)
            ws.cell(row=row, column=5, value=track.mainArtist)
            ws.cell(row=row, column=6, value=track.featArtist)
            ws.cell(row=row, column=8, value=track.profile)
            ws.cell(row=row, column=9, value=track.composer)
            ws.cell(row=row, column=10, value=track.writer)
            ws.cell(row=row, column=11, value=track.producer)
            ws.cell(row=row, column=12, value=track.genre if track.genre else payload.genre)
            ws.cell(row=row, column=13, value=track.lang)
            ws.cell(row=row, column=14, value=track.year)
            ws.cell(row=row, column=15, value="exclusively licensed to MMusic Records")
            ws.cell(row=row, column=16, value="exclusively licensed to MMusic Records")
            ws.cell(row=row, column=19, value=track.explicit)
            ws.cell(row=row, column=20, value=track.tiktok)
            ws.cell(row=row, column=21, value=track.lyrics)
        
        temp_filename = f"/tmp/{release_id}_label_form.xlsx"
        wb.save(temp_filename)
        
        creds = drive_service.get_credentials()
        service = drive_service.get_drive_service(creds)
        
        query = f"'{target_folder_id}' in parents and trashed=false"
        results = service.files().list(q=query, spaces='drive', fields='files(id, name)').execute()
        items = results.get('files', [])
        
        for item in items:
            try:
                service.files().delete(fileId=item['id']).execute()
            except:
                pass
                
        new_name = f"[MMusic Records] Label Form - {payload.productTitle}.xlsx"
        uploaded = drive_service.upload_file(service, temp_filename, new_name, target_folder_id)
        
        os.remove(temp_filename)
        update_release_status(release_id, "label_form")

        # Lưu metadata label form vào Firestore để Admin trích xuất email
        if db:
            doc_ref = db.collection('releases').document(release_id)
            doc_ref.update({
                "label_form_data": {
                    "productTitle": payload.productTitle,
                    "productType": payload.productType,
                    "productArtist": payload.productArtist,
                    "genre": payload.genre,
                    "releaseDate": payload.releaseDate,
                    "releaseTime": payload.releaseTime,
                    "explicit": payload.tracks[0].explicit if payload.tracks else "No",
                },
                "label_form_link": uploaded.get("webViewLink", "")
            })

        return JSONResponse({"message": "Success", "link": uploaded.get("webViewLink")})
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JSONResponse({"error": str(e)}, status_code=500)

@app.post("/api/release/{release_id}/submit-contract")
async def submit_contract(release_id: str, request: Request):
    release_data = get_release(release_id)
    if not release_data:
        return JSONResponse({"error": "Release not found"}, status_code=404)
        
    folder_ids = release_data.get("folder_ids", {})
    target_folder_id = folder_ids.get("5. Contract Info")
    
    if not target_folder_id:
        return JSONResponse({"error": "Contract Info folder not found. Please re-run automation."}, status_code=500)
    
    try:
        from docx import Document
        from docx.shared import Inches
        import uuid
        
        form_data = await request.form()
        contract_type = form_data.get("contractType", "personal")
        artist_name = release_data.get("artist_name", "Artist")
        
        doc = Document()
        doc.add_heading("THÔNG TIN HỢP ĐỒNG PHÁT HÀNH NHẠC", 0)
        
        if contract_type == "personal":
            doc.add_heading("I. Thông tin khách hàng (Cá nhân)", level=1)
            fields = [
                ("Họ tên", "p_name"), ("Nghệ danh", "p_stagename"), ("Sinh ngày", "p_dob"), 
                ("CCCD số", "p_cccd"), ("Cấp ngày", "p_cccd_date"), ("Tại", "p_cccd_place"),
                ("Địa chỉ cũ", "p_address_old"), ("Địa chỉ hiện tại", "p_address_new"), 
                ("Điện thoại", "p_phone"), ("Email", "p_email"), ("Mã số thuế", "p_tax"),
                ("Số tài khoản", "p_bank_num"), ("Tên tài khoản", "p_bank_name"), 
                ("Ngân hàng", "p_bank_brand"), ("Chi nhánh", "p_bank_branch")
            ]
            for label, key in fields:
                val = form_data.get(key, '')
                if val: doc.add_paragraph(f"{label}: {val}")
                
            doc.add_heading("II. Thông tin bản ghi", level=1)
            r_fields = [
                ("Tên Bản ghi", "r_title"), ("Tác giả", "r_author"), 
                ("Người trình bày", "r_singer"), ("Nhà Sản xuất", "r_producer"), 
                ("Chủ sở hữu", "r_owner")
            ]
            for label, key in r_fields:
                val = form_data.get(key, '')
                if val: doc.add_paragraph(f"{label}: {val}")
                
        else:
            doc.add_heading("I. Thông tin khách hàng (Công ty)", level=1)
            fields = [
                ("CÔNG TY", "c_name"), ("Đại diện", "c_rep"), ("Chức vụ", "c_role"), 
                ("Địa chỉ", "c_address"), ("Email", "c_email"), ("Số điện thoại", "c_phone"),
                ("Tài khoản", "c_bank_num"), ("CHỦ TK", "c_bank_name"), 
                ("Tại ngân hàng", "c_bank_brand"), ("Mã số thuế", "c_tax")
            ]
            for label, key in fields:
                val = form_data.get(key, '')
                if val: doc.add_paragraph(f"{label}: {val}")

            doc.add_heading("II. Thông tin bản ghi", level=1)
            r_fields = [
                ("Tên Bản ghi", "r_title"), ("Tác giả", "r_author"), 
                ("Người trình bày", "r_singer"), ("Nhà Sản xuất", "r_producer"), 
                ("Chủ sở hữu", "r_owner")
            ]
            for label, key in r_fields:
                val = form_data.get(key, '')
                if val: doc.add_paragraph(f"{label}: {val}")
        
        doc.add_paragraph("")
        doc.add_paragraph("Thời hạn hợp tác (MMusic điền): " + "_" * 40)
        doc.add_paragraph("Tỷ lệ chia sẻ doanh thu: 70/30 (nghệ sĩ 70%)")
        doc.add_paragraph("")
        
        doc.add_heading("III. Hồ sơ đính kèm", level=1)
        
        creds = drive_service.get_credentials()
        service = drive_service.get_drive_service(creds)
        
        file_keys = ["cccd_front", "cccd_back", "vneid_img", "c_dkkd", "c_uyquyen"]
        files_to_clean = []
        
        for key in file_keys:
            uploaded_file = form_data.get(key)
            if uploaded_file and hasattr(uploaded_file, "filename") and uploaded_file.filename:
                ext = os.path.splitext(uploaded_file.filename)[1].lower()
                tmp_path = f"/tmp/{uuid.uuid4()}{ext}"
                
                contents = await uploaded_file.read()
                with open(tmp_path, "wb") as f:
                    f.write(contents)
                files_to_clean.append(tmp_path)
                
                doc.add_paragraph(f"--- Đính kèm ({key}): {uploaded_file.filename} ---")
                if ext in ['.jpg', '.jpeg', '.png', '.gif']:
                    try:
                        doc.add_picture(tmp_path, width=Inches(6.0))
                    except Exception as e:
                        doc.add_paragraph(f"(Lỗi chèn ảnh vào Word: {e})")
                else:
                    doc.add_paragraph("(File này không phải ảnh, đã được upload riêng lên thư mục Contract Info)")
                    drive_service.upload_file(service, tmp_path, uploaded_file.filename, target_folder_id)

        clean_name = "".join([c if c.isalnum() else "_" for c in str(artist_name)])
        docx_filename = f"Contract_Info_{clean_name}_{release_id}.docx"
        docx_path = f"/tmp/{docx_filename}"
        doc.save(docx_path)
        
        # Override the old file if exists
        query = f"'{target_folder_id}' in parents and name contains 'Contract_Info_' and trashed=false"
        results = service.files().list(q=query, spaces='drive', fields='files(id)').execute()
        for item in results.get('files', []):
            try: service.files().delete(fileId=item['id']).execute()
            except: pass

        drive_service.upload_file(service, docx_path, docx_filename, target_folder_id)
        
        if os.path.exists(docx_path):
            os.remove(docx_path)
        for f in files_to_clean:
            if os.path.exists(f):
                os.remove(f)
                
        update_release_status(release_id, "contract")
        return JSONResponse({"message": "Successfully generated and uploaded Contract Info."})
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JSONResponse({"error": str(e)}, status_code=500)

@app.post("/api/release/{release_id}/submit-pr-form")
async def submit_pr_form(release_id: str, payload: PRFormRequest):
    release_data = get_release(release_id)
    if not release_data:
        return JSONResponse({"error": "Release not found"}, status_code=404)
        
    folder_ids = release_data.get("folder_ids", {})
    target_folder_id = folder_ids.get("product_root")
    
    if not target_folder_id:
        return JSONResponse({"error": "Product root folder ID not found"}, status_code=500)
        
    try:
        from docx import Document
        import uuid
        
        doc = Document()
        doc.add_heading("BẢNG CÂU HỎI TRUYỀN THÔNG (PR QUESTIONNAIRE)", 0)
        
        questions = [
            "1. Mô tả một vài dòng về bài hát",
            "2. Giới thiệu về bản thân và về sản phẩm bạn sắp phát hành",
            "3. Bạn có đến từ hội nhóm, crew hay team nào không? Có thể giới thiệu qua một chút không?",
            "4. Cảm hứng sáng tác bài hát này của bạn đến từ đâu?",
            "5. Nếu là sản phẩm hợp tác, có thể chia sẻ một chút về lí do hợp tác và sự đồng điệu như thế nào?",
            "6. Thời điểm phát hành bài hát này có gì đặc biệt hay không?",
            "7. Có ý nghĩa gì sâu xa từ cách đặt tên bài hát của bạn không?",
            "8. Chất liệu âm nhạc chính/Thể loại của bài hát là gì? Vì sao?",
            "9. Bạn mong muốn gửi gắm điều gì tới khán giả khi nghe bài hát?",
            "10. Sản phẩm này là một trải nghiệm mới hay là thế mạnh trước giờ của bạn?",
            "11. Bạn có kỷ niệm đặc biệt nào trong quá trình sản xuất bài hát không?",
            "12. Bạn có thể chia sẻ thêm một chút về những dự định trong tương lai gần hay không?",
            "13. Chi tiết thêm nghệ sĩ muốn chia sẻ:"
        ]
        
        answers = [
            payload.q1, payload.q2, payload.q3, payload.q4, payload.q5,
            payload.q6, payload.q7, payload.q8, payload.q9, payload.q10,
            payload.q11, payload.q12, payload.q13
        ]
        
        for q, a in zip(questions, answers):
            doc.add_heading(q, level=2)
            doc.add_paragraph(a if a.strip() else "(Không có chia sẻ)")
            
        artist_name = release_data.get("artist_name", "Artist")
        product_title = release_data.get("product_name", "Product")
        clean_product = "".join([c if c.isalnum() else "_" for c in str(product_title)])
        
        docx_filename = f"[PR] {clean_product}_{release_id}.docx"
        docx_path = f"/tmp/{docx_filename}"
        doc.save(docx_path)
        
        creds = drive_service.get_credentials()
        service = drive_service.get_drive_service(creds)
        drive_service.upload_file(service, docx_path, docx_filename, target_folder_id)
        
        if os.path.exists(docx_path):
            os.remove(docx_path)
                
        update_release_status(release_id, "pr_form")
        return JSONResponse({"message": "Successfully generated and uploaded PR Questionnaire."})
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JSONResponse({"error": str(e)}, status_code=500)


@app.post("/api/release/{release_id}/upload")
async def upload_release_files(
    release_id: str,
    category: str = Form(...),
    files: List[UploadFile] = File(...)
):
    release_data = get_release(release_id)
    if not release_data:
        return JSONResponse({"error": "Release not found"}, status_code=404)
    
    folder_ids = release_data.get("folder_ids", {})
    
    category_map = {
        "label_form": "1. Label form (track info)",
        "audio": "2. Audio (16-24 BIT DEPTH)",
        "artwork": "3. Artwork (3000x3000)",
        "artist_picture": "4. Artist_s Picture (cover - profile)",
        "contract": "5. Contract Info",
        "mv": "6. MV (nếu có)",
        "canvas": "7. Spotify Canvas (4-8s, 9_16)",
        "others": "8. Other materials"
    }
    
    target_folder_name = category_map.get(category)
    if not target_folder_name:
        return JSONResponse({"error": "Invalid category"}, status_code=400)
    
    target_folder_id = folder_ids.get(target_folder_name)
    if not target_folder_id:
        return JSONResponse({"error": "Folder ID not found for this category"}, status_code=500)

    try:
        creds = drive_service.get_credentials()
        service = drive_service.get_drive_service(creds)
        uploaded_results = []
        for file in files:
            temp_path = f"/tmp/{uuid.uuid4()}_{file.filename}"
            with open(temp_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
            
            uploaded = drive_service.upload_file(
                service, 
                temp_path, 
                file.filename, 
                target_folder_id
            )
            os.remove(temp_path)
            uploaded_results.append({
                "filename": file.filename, 
                "link": uploaded.get("webViewLink")
            })
            
        update_release_status(release_id, category)
        return JSONResponse({"message": "Upload successful", "files": uploaded_results})
    except Exception as e:
        print(f"Upload error: {e}")
        return JSONResponse({"error": str(e)}, status_code=500)
